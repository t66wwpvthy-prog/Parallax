from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BRASS = RGBColor(0xB0, 0x7D, 0x34)
INK   = RGBColor(0x1d, 0x19, 0x13)
GREY  = RGBColor(0x6b, 0x62, 0x53)
ITAL  = RGBColor(0x8a, 0x7c, 0x61)

doc = Document()
for s in doc.sections:
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)

# Title
t = doc.add_paragraph()
r = t.add_run("PARALLAX — Master Idea Bank"); r.bold = True
r.font.size = Pt(22); r.font.color.rgb = INK
sub = doc.add_paragraph("Everything parked, consolidated. Retirement-planning simulator for advisors.")
sub.runs[0].font.size = Pt(11); sub.runs[0].font.italic = True; sub.runs[0].font.color.rgb = GREY
meta = doc.add_paragraph("Generated 2026-06-03 · sources: ROADMAP.md + NOTES.md + CLAUDE.md · tags: NEXT · BIG (own session) · PARKED · SKIP · done")
meta.runs[0].font.size = Pt(8.5); meta.runs[0].font.color.rgb = RGBColor(0x9a,0x8f,0x7c)
doc.add_paragraph()

def section(title):
    h = doc.add_heading(level=1)
    run = h.add_run(title); run.font.color.rgb = BRASS; run.font.size = Pt(14); run.bold = True

def lede(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(9.5); p.runs[0].font.italic = True; p.runs[0].font.color.rgb = GREY

def bullet(lead, rest="", tag=""):
    p = doc.add_paragraph(style="List Bullet")
    a = p.add_run(lead); a.bold = True; a.font.size = Pt(10.5); a.font.color.rgb = INK
    if rest:
        b = p.add_run(" — " + rest); b.font.size = Pt(10.5); b.font.color.rgb = RGBColor(0x33,0x2f,0x29)
    if tag:
        g = p.add_run("  [" + tag + "]"); g.font.size = Pt(8); g.bold = True; g.font.color.rgb = ITAL

section("1 · Engine — material gaps  (truth source)")
lede("The single truth source. The UI never adds math — only views/levers of what the engine emits.")
bullet("Real assets in engine", "homes / property / business → “sell X to fund Y.”", "BIG")
bullet("Tax-planning engine", "Roth conversions, bracket management, IRMAA, RMDs, contribution-side deduction. Its own engine.", "BIG")
bullet("Survivor SS benefits", "widow step-up (~100% of deceased's); major couples lever, unmodeled.")
bullet("Healthcare scales independently of lifestyle", "spendMult currently scales medical too; skews Solve-For.")
bullet("LTC cost escalation", "~3–5%/yr above CPI; currently flat real.")
bullet("One-time inflow support", "relocation / downsizing; current lumpSum is outflow-only.")
bullet("Map the other engines", "tax, estate… before building piecemeal.")
bullet("Cleanup", "ASSET_STATS[k].mean computed but unused — deletable.")

section("2 · Core differentiators — the moat  (BIG)")
lede("The defensible layer: deterministic, explainable, multi-variable decision-solving — not “Monte Carlo + sliders.”")
bullet("Solve-For mode", "named life questions: ① “Can I retire earlier?” (build first) ② “Can we afford this big goal?” ③ (later) “Spend more?” On-demand (one engine job/click); output = a new scenario column; exactly one free lever.", "NEXT-ish")
bullet("Explainable constraint solver", "multi-variable search returning a small Pareto set of least-disruptive paths; auditable replay. Also fixes the solo-vs-combined display confusion.")
bullet("Decision Surface heatmap", "retire-age × spend, confidence contours, saved scenarios as clickable points. Best embodiment of “Parallax models interactions.” Alternate view in Scenarios.")

section("3 · High-value views of existing truth  (low lift)")
lede("Mostly a VIEW of data the sims already emit — no new math.")
bullet("Failure Anatomy", "click a % ring → depletion-age histogram, earliest pressure year, success/depleted split, one representative failed path in the cash-flow drawer. Best ROI.")
bullet("Scenario Receipt", "per-scenario delta drawer (levers moved + confidence/median before→after + worst market). Meeting takeaway / export. Data, not prose.")
bullet("Sequence Pressure Strip", "first ~10 retirement years as cells (return / withdrawal / ending balance), cursor synced to chart + drawer. Explains why 1973 / 2000 hurt.")
bullet("Funding Bridge / bridge-years", "stacked age timeline of income layers; the pre-SS/pension gap.")
bullet("Honest derived tiles", "longevity buffer, max sustainable withdrawal %, “could retire at X.”")
bullet("Confidence-by-year bars", "+ a single “Confidence Age” — only if it's real per-year data.")

section("4 · Scenarios / input model")
bullet("Account-types picker", "“+ add investment account” with type picker (401k/403b/457/SEP/SIMPLE/solo-401k → pre-tax).", "mock built, awaiting OK")
bullet("Scenarios “shared-track” redesign", "kill the 3-column repetition; one shared track per lever, dots split only on divergence. Approved-in-spirit.", "BIG")
bullet("Pension claim-age analysis", "sub-mode: claim-age sweep (62→70) + breakeven.")
bullet("Real-assets input", "Money Pro balance-details layout (pairs with engine real-assets).")
bullet("“Duplicate scenario” action", "distinct from Add.")

section("5 · Sequencing")
bullet("Strategy Fork on one path", "fix ONE market, compare strategies on it. Your #1 advisor function. Decide Scenarios-vs-Sequencing home first (overlap).")
bullet("Recovery Tunnel", "full years-underwater valley view.")
bullet("Sequence Duration / damage-window labels", "“short shock” (’08–09, V-shaped) vs “long grind” (’66 / ’73). A label, not new math.")

section("6 · This session's adds  (2026-06-03)")
bullet("Home sale event + cap-gains on sale", "proceeds = value − mortgage − cap-gains tax; needs the $250k/$500k-MFJ primary-residence exclusion + a capitalImprovements basis field. purchasePrice already captured & inert. Extends “real assets.”")
bullet("taxClass on income streams", "ordinary / tax-free / cap-gains; today all other income is flat ordinary. An engine field, not a label.")
bullet("Sustainable-withdrawal solve", "back into a safe annual portfolio draw, framed as “sustainable withdrawal,” not income.")
bullet("Income-stream COLA", "variable income is flat-real today; let fixed annuities erode.")
bullet("Rolling-period analysis", "extended block bootstrap — sweep every real contiguous historical window, not just the named Sequencing years.")
bullet("Goals page — interesting layout", "not the ledger style. End-date windows already built; visual redesign pinned.")

section("7 · Future surfaces / views")
bullet("Comprehensive plan view (read-only)", "the whole plan as a presentation-grade client deliverable. Edits stay on input pages.")
bullet("Prospect-level view", "a lighter “quick illustration” tier for prospects/sales (a reduced view of the same engine).")
bullet("Investment comparisons", "proposed-vs-current / efficient-frontier. ⚠ overlaps Scenarios — pin the distinct angle first.")

section("8 · Smaller wins / hygiene")
bullet("Rename “Net Worth” tab → “Plan” / “Household”", "now holds cash flow / goals / snapshot. Quick win.")
bullet("Codify visual grammar", "dashed underline = editable; solid = result; brass = selected/caution; clay = pressure; teal = baseline. Write it down.")
bullet("Assumption ledger", "Entered / Derived / Simulated / Historical tags. ⚠ must NOT become the rejected “needs-review” machinery.")

section("9 · Open decisions")
bullet("Demo mortgage", "keep (baseline 84.9%) vs strip to a paid-off house (~86.5%)?")
bullet("Net Worth layout", "make Assets the wider column (Liabilities is often sparse)?")
bullet("SS framing", "lifetime-dollars trade-off vs the % circle (weak % mover for high-spend households).")
bullet("Solve-For display", "solo-vs-combined so an advisor can't misread the column %.")
bullet("Liabilities", "base-plan input vs per-scenario construct.")
bullet("Strategy Fork", "Sequencing feature or Scenarios feature?")

section("10 · Skip / do-not-relitigate  (logged on purpose)")
bullet("Cash-flow River / Sankey", "decorative, no new interaction.")
bullet("Resilience Matrix as a new tab", "old Stress Test reincarnated; salvage only the click-to-clone link.")
bullet("Rebrand away from “Parallax”", "business decision, not engineering.")
bullet("Reversed-sequence view", "killed (a timeline that never happened).")
bullet("Allocation as a solve lever", "re-add only with deliberate framing.")

doc.add_paragraph()
note = doc.add_paragraph()
nr = note.add_run("Standing doctrine.  "); nr.bold = True; nr.font.size = Pt(9.5); nr.font.color.rgb = BRASS
nb = note.add_run("Make requirements less dumb → delete before adding → simplify → cut beats refactor. The engine is the single truth source; the UI only changes views/levers, never the math. Green is allowed — no mint. Mock visual changes first. Verify by looking. Push to both the working branch and main.")
nb.font.size = Pt(9.5); nb.font.color.rgb = GREY; nb.font.italic = True

doc.save("/home/user/Parallax/Parallax-Idea-Bank.docx")
print("Done.")
