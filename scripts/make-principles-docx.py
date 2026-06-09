from docx import Document
from docx.shared import Pt, RGBColor, Inches

BRASS=RGBColor(0xB0,0x7D,0x34); INK=RGBColor(0x1d,0x19,0x13); GREY=RGBColor(0x6b,0x62,0x53)

doc=Document()
for s in doc.sections:
    s.top_margin=Inches(0.9); s.bottom_margin=Inches(0.9); s.left_margin=Inches(1.0); s.right_margin=Inches(1.0)

t=doc.add_paragraph(); r=t.add_run("PARALLAX — Principles & Jewels"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=INK
sub=doc.add_paragraph("What the tool is, how we build it, and the lessons we paid for."); sub.runs[0].font.size=Pt(11); sub.runs[0].font.italic=True; sub.runs[0].font.color.rgb=GREY
doc.add_paragraph()

def section(title):
    h=doc.add_heading(level=1); run=h.add_run(title); run.font.color.rgb=BRASS; run.font.size=Pt(14); run.bold=True

def bullet(lead, rest=""):
    p=doc.add_paragraph(style="List Bullet")
    a=p.add_run(lead); a.bold=True; a.font.size=Pt(10.5); a.font.color.rgb=INK
    if rest:
        b=p.add_run(" — "+rest); b.font.size=Pt(10.5); b.font.color.rgb=RGBColor(0x33,0x2f,0x29)

section("The 3 core principles")
bullet("Show the interaction.", "the value is watching one lever move another and the client SEEING it — not isolated stats. Every view must show an interaction or a truth, or it shouldn't exist.")
bullet("Show the truth, as close as we can get.", "the program shows the story; the advisor tells it. If the picture is right, the sentence is redundant — cut it.")
bullet("Stay neutral.", "the tool has no opinion. As willing to say “you're fine, spend more” as “this won't last.” Showing good news isn't optimism — it's removing the fear bias every other tool is built on. That neutrality is WHY an advisor can trust it in front of a client.")

section("How we build")
bullet("Make the requirement less dumb.", "question it; trace back to fundamentals.")
bullet("Subtract before adding.", "deleting is the default. Ask what removing would cost.")
bullet("Simplify.", "fewer steps, fewer abstractions. Fewer rules = less code.")
bullet("The engine is the one source of truth.", "the screens only show or adjust it, never invent math. Don't touch the engine without explicit agreement.")
bullet("When two builds compete,", "pick the one most faithful to the truth.")
bullet("Don't let it bloat.", "the last build became “a bloated PowerPoint of disconnected concepts.” When tempted to add scope, say so instead of doing it.")

section("The jewels (learned the hard way)")
bullet("Logic checks lie; pixels don't.", "look at the rendered screen before calling it done. (The cash-flow drawer shipped 2px tall for 10 messages because nobody looked.)")
bullet("Faithful-and-ugly beats clever-and-broken.", "a clever CSS shortcut silently wrecked the whole look.")
bullet("Mock big visual changes first.", "static study → screenshot → you approve → then build.")
bullet("Terminal wealth is NOT the goal.", "it's only a ranking device. Plan for spending security, surviving bad markets, and meeting goals.")
bullet("Look rules.", "no bright white, no mint green, no code fonts. Use only the currently approved theme tokens for the active build. Charts smooth and trackable, never jagged.")

section("How we work together")
bullet("Short over long.", "lead with the result; skip the recap.")
bullet("Never guess.", "verify (run it, read it) or say “I can't verify that.”")
bullet("One feature ≈ one session.", "finish it, bank only current truth, then clear — stale notes must never compete with doctrine.")

doc.save("/home/user/Parallax/Parallax-Principles.docx"); print("Done.")