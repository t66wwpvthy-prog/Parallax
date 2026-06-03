from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("PARALLAX")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2c, 0x30, 0x38)

subtitle = doc.add_paragraph("Six things the engine does that most planning tools don't.")
subtitle.runs[0].font.size = Pt(11)
subtitle.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
subtitle.runs[0].font.italic = True

doc.add_paragraph()

points = [
    (
        "01 — Real returns. Inflation is native, not bolted on.",
        "Every return in the dataset is already inflation-adjusted. The engine works entirely in today's dollars — so when it says your portfolio grows, that's real purchasing power, not nominal noise.",
        "Most tools model nominal returns and subtract inflation separately, which compounds the error across decades."
    ),
    (
        "02 — Withdrawals happen throughout the year, not at year-end.",
        "Spending is spread across 12 monthly draws. The account keeps earning returns while money is still in it. In a bad year, you're not losing a full year of growth before the first dollar leaves.",
        "Year-end withdrawal models overstate the damage in down years and understate compounding in good ones."
    ),
    (
        "03 — Returns are sampled in blocks, not drawn randomly year by year.",
        "The Monte Carlo draws multi-year blocks from 98 years of real return data (1928–2025). A crash year drags its neighbors. That preserves the clustering of bad sequences — the thing that actually ends retirement plans.",
        "Independent year-by-year sampling never produces the sustained bad runs that sequence-of-returns risk actually looks like."
    ),
    (
        "04 — Every scenario runs the same market draws.",
        "When you compare two scenarios, both live through the identical 1,000 market sequences. Any difference in the outcome is purely the decision you changed — not sampling luck.",
        "Most tools run independent simulations per scenario, so you can't know if the number moved because of your lever or noise in the draw."
    ),
    (
        "05 — Three account types. Withdrawals are sequenced, not blended.",
        "Money comes from taxable first, then traditional (RMDs kick in at 73), then Roth. The engine tracks cost basis in taxable accounts and treats traditional withdrawals as ordinary income. Tax drag is modeled in the order it actually happens.",
        "A blended tax rate applied to a single portfolio misses the sequence entirely — and the sequence is what determines the tax bill."
    ),
    (
        "06 — Historical paths are the real thing.",
        "The Sequencing view runs the actual 1973 return sequence, year by year, in order. Not 'a representative bad period' — that period. 1929, 1966, 1973, 2000, 2008. Your plan, living through each one.",
        "Stylized stress tests (e.g. 'a 30% drop') are arbitrary. The historical record is the most honest stress test available."
    ),
]

for heading, body, contrast in points:
    h = doc.add_paragraph()
    run = h.add_run(heading)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2c, 0x30, 0x38)

    b = doc.add_paragraph(body)
    b.runs[0].font.size = Pt(11)
    b.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    c = doc.add_paragraph(contrast)
    c.runs[0].font.size = Pt(10)
    c.runs[0].font.italic = True
    c.runs[0].font.color.rgb = RGBColor(0x9b, 0x88, 0x55)

    doc.add_paragraph()

doc.save("/home/user/Parallax/parallax-engine-explainer.docx")
print("Done.")
